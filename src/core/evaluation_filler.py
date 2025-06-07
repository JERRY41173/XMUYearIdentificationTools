# -*- coding: utf-8 -*-
"""
评语填写核心模块
功能：自动填写学年鉴定表中的各种评语
"""

import os
import random
import shutil
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


class EvaluationFiller:
    """评语填写器"""
    
    def __init__(self):
        self.academic_years = ["2021-2022", "2022-2023", "2023-2024", "2024-2025"]
        self.total_expected_evaluations = 7  # 4个学年意见 + 3个综合鉴定表评语
        
        # 初始化评语库
        self._init_evaluation_templates()
    
    def _init_evaluation_templates(self):
        """初始化评语模板"""
        # 学年意见
        self.academic_year_opinions = [
            [
                "    该生在学期间，思想态度端正，遵守校规校纪，学习刻苦，成绩良好，同时能积极参与实践活动，团队协作能力强。该生综合素质高，是一名品学兼优的大学生。",
                "    该生在学期间，思想态度端正，积极向上，展现出较强的责任感。学习认真，富有创新精神。团队协作能力强，与同学相处较好。该生全面发展，是一名品学兼优的大学生。"
            ],
            [
                "    该生在学期间积极向上，学习努力，并能积极参与各类活动，不断提升自己的综合素质，日常生活中尊重师长，能与同学和睦相处。该生是一名表现良好的大学生，具有较大的发展潜力。",
                "    该生在学期间学习努力，成绩良好，能够认真听讲、积极思考，具有良好的个人修养。同时该生也积极参与集体活动，为班级和学院争光。该生表现良好，具有较大的发展潜力。"
            ],
            [
                "    该生在学期间展现出了良好的综合素质，学习认真刻苦，还注重培养自己的兴趣爱好和特长，积极参与各类活动和比赛，是一名品学兼优的大学生。",
                "    该生在学期间学习认真刻苦，在课堂上表现出色，课后也能主动寻求知识，不断提升学术水平；此外也注重培养自己的兴趣爱好和特长，集体生活中也能够积极融入团队，是一名品学兼优的大学生。"
            ],
            [
                "    该生在学期间思想态度端正，政治立场正确；学习认真，兴趣爱好广泛，积极参与院校各类活动，表现良好；尊敬师长，团结同学，人际关系良好，是一名优秀的毕业生。",
                "    该生在学期间严格遵守校纪校规，政治立场正确；学习认真，具备良好的理论知识和实践能力；兴趣爱好广泛，工作认真，具有一定的组织协调和管理能力；尊敬师长，团结同学，是一名优秀的毕业生。"
            ],
            [
                "    该生在学期间遵纪守法，思想态度端正，积极向上。学业成绩突出，专业知识扎实，展现出良好的学术素养。在课外活动中，该生展现出良好的领导力和组织协调能力，深受师生认可。同时，该生具备良好的沟通能力和团队协作精神，是一名全面发展的优秀毕业生。"
            ]
        ]
        
        # 班团组织鉴定
        self.class_organization_evaluations = [
            "\n    该生思想积极要求进步，专业学习认真踏实，平时能够严格要求自己，能很好地遵守学校纪律，集体荣誉感强，参加班集体活动，能与同学友好相处，尊敬师长，是一位综合素质全面的大学生。",
            "\n    该生能很好地遵守学校纪律，集体荣誉感强，积极参与各项社会实践与集体活动，关心集体，热心为班级服务，积极肯干，尊敬师长，团结同学；学习上刻苦努力，专业基础扎实，成绩优秀；在任学生干部期间，表现出较强的组织能力，热心为同学服务，是一位有理想、有抱负、全面发展的大学生。",
            "\n    该生学习态度端正，动手能力强，积极参加各类科创活动，具备学习的自觉性和主动新，在科研方面具有有创新意识，有独立分析问题、解决问题的能力，尊敬老师，能与同学友好相处，是一个品学兼优的好学生。",
            "\n    该生积极向上，活泼开朗，能与同学友好相处，助人为乐，积极参加班级的活动，集体荣誉感强，热爱班级，多次参加各类学校、学院活动，学习认真，有钻研精神和创新意识，各方面表现都比较优秀。",
            "\n    该生能很好地履行大学生行为规范。在学习上肯下功夫, 严格要求自己，勤学好问,努力钻研,尊敬师长,团结同学,积极参加各类集体活动和社会实践活动。学习目标明确,刻苦认真,是位综合素质较高的学生。"
        ]
        
        # 班主任综合评语
        self.class_teacher_evaluations = [
            "\n    在校期间，该生思想上进，作风严谨，思维活跃，学习上踏实认真，能从各方面严格要求自己；性格开朗，乐观向上，积极参加各种志愿服务活动；平时待人随和而友善，团结同学，尊敬师长，遵守校规校纪，乐于助人；是一名素质优良的大学生。",
            "\n    该生能够坚持四项基本原则，拥护党的领导，遵守各项法律法规和学校规定的各项规章制度；学习上，认真刻苦，具有扎实的专业基础知识，较强的学习和创新能力，学习成绩优异；工作认真负责，积极主动，尽职尽责，为老师和同学做大量工作；性格开朗、稳重、有活力，待人热情，真诚，具有良好的团队合作精神和沟通能力，是一名品学兼优、德才兼备的大学生。",
            "\n    该生思想上进，作风严谨，积极参加学校和班级组织的各项工作；学习态度端正，积极参加各项科创活动，并取得优异成绩；对待同学，真诚热心，乐于助人；动手能力和自学能力较强，具备良好的分析问题和独立思考的能力，并且有很强的集体荣誉感；勇于面对困难，敢于迎接挑战；是一名品学兼优的大学生。",
            "\n    该生在思想道德方面始终坚持正确的价值观和道德观，遵守校规校纪，尊重师长，团结同学。关心集体，乐于助人，总是能够在同学需要帮助时伸出援手。在学习上始终保持着严谨的学习态度和扎实的专业知识基础。课堂上与老师和同学进行深入讨论，不断拓宽自己的知识视野。同时，注重将所学知识运用到实际生活中，在班团活动中表现出了强烈的集体荣誉感和责任感。综上所述，该生是一位全面发展、综合素质较高的优秀同学。",
            "\n    该生在思想道德方面，始终坚持正确的价值观和道德观，遵守校规校纪，尊重师长，团结同学。能够在困难和挫折面前保持积极乐观的态度，该生拥有较出色的学习能力和扎实的基础。经常参与学术活动，不断拓展自己的学术视野。该生展现出高度的责任感和团队协作精神。善于沟通，乐于助人，综述，该生是一位全面发展的优秀学生。"
        ]
        
        # 学院意见
        self.college_opinions = [
            "\n    该生在学期间，思想态度端正，积极向上，遵守校规校纪，自我约束力强，展现出较强的责任感。学习刻苦，成绩良好，科研认真严谨，富有创新精神。同时，积极参与实践活动，团队协作能力强。该生全面发展，综合素质高，是一名品学兼优的大学生。",
            "\n    该生在学期间积极向上，学习努力，成绩良好，能够认真听讲、积极思考，并在课后及时复习巩固。该生积极参与各类实践活动，不断提升自己的综合素质。遵守校规校纪，尊重师长，与同学和睦相处，展现了良好的个人修养。同时，该生也积极参与集体活动，为班级和学院争光。总之，该生是一名表现良好的大学生，具有较大的发展潜力。",
            "\n    该生在学期间展现出了良好的综合素质。学习认真刻苦，在课堂上表现出色，在课后也能主动寻求知识，不断提升自己的学术水平。此外，该生还注重培养自己的兴趣爱好和特长，积极参与各类社团活动和比赛，取得了一定成绩。在集体生活中，该生能够积极融入团队，与同学们友好相处，共同为集体荣誉而努力。该生遵守校规校纪，是一名品学兼优的大学生。",
            "\n    该生在学期间严格遵守校纪校规，思想态度端正，政治立场正确。学习认真，具备良好的理论知识和实践能力；兴趣爱好广泛，积极参与院校各类活动，表现良好；做事稳重踏实，具有一定的组织协调和管理能力；尊敬师长，团结同学，人际关系良好。总之，是一名优秀的毕业生。",
            "\n    该生在学期间遵纪守法，思想态度端正，积极向上。学业成绩突出，专业知识扎实，展现出良好的学术素养。在课外活动中，该生展现出良好的领导力和组织协调能力，深受师生认可。同时，该生具备良好的沟通能力和团队协作精神，是一名全面发展的优秀毕业生。"
        ]
    
    def process_folder(self, folder_path):
        """处理文件夹中的所有docx文件"""
        if not os.path.exists(folder_path):
            raise Exception(f"文件夹不存在: {folder_path}")
        
        # 创建错误文件夹
        error_folder = os.path.join(folder_path, "error")
        
        # 统计变量
        total_files = 0
        success_files = 0
        error_files = 0
        
        # 获取所有docx文件
        docx_files = [f for f in os.listdir(folder_path) 
                      if f.endswith('.docx') and not f.startswith('~') and f != "error"]
        
        if not docx_files:
            raise Exception("未找到任何docx文件")
        
        # 处理每个文件
        for filename in docx_files:
            total_files += 1
            file_path = os.path.join(folder_path, filename)
            
            # 处理文件
            success = self._process_single_file(file_path, error_folder)
            if success:
                success_files += 1
            else:
                error_files += 1
        
        return {
            'total': total_files,
            'success': success_files,
            'error': error_files,
            'error_folder': error_folder if error_files > 0 else None
        }
    
    def _process_single_file(self, file_path, error_folder):
        """处理单个docx文件，自动填写评语"""
        try:
            # 打开文档
            doc = Document(file_path)
            
            # 初始化计数器
            academic_years_filled = 0
            comprehensive_filled = 0
            
            # 处理各学年的学院意见
            for i, year_suffix in enumerate(self.academic_years):
                try:
                    table = self._find_academic_year_table(doc, year_suffix)
                    if table:
                        # 选择评语
                        if i < len(self.academic_year_opinions) and self.academic_year_opinions[i]:
                            evaluation_text = random.choice(self.academic_year_opinions[i])
                        else:
                            evaluation_text = "学生在本学年表现良好。"
                        
                        if self._fill_evaluation_text(table, "学院 意见", evaluation_text):
                            academic_years_filled += 1
                except Exception:
                    pass
            
            # 处理综合鉴定表（表格12）
            try:
                if len(doc.tables) >= 12:
                    table12 = doc.tables[11]  # 索引从0开始
                    
                    # 1. 班团组织鉴定 (第4行)
                    if self._fill_table_cell(table12, 3, 1, random.choice(self.class_organization_evaluations)):
                        comprehensive_filled += 1
                    
                    # 2. 班主任综合评语 (第11行)
                    if self._fill_table_cell(table12, 10, 1, random.choice(self.class_teacher_evaluations)):
                        comprehensive_filled += 1
                    
                    # 3. 学院意见 (第17行)
                    if self._fill_table_cell(table12, 16, 1, random.choice(self.college_opinions)):
                        comprehensive_filled += 1
            except Exception:
                pass
            
            # 计算总数
            total_filled = academic_years_filled + comprehensive_filled
            
            # 严格检查：只有所有7个评语都成功填写才算成功
            if total_filled == self.total_expected_evaluations:
                # 保存文档
                doc.save(file_path)
                return True
            else:
                # 将文件移动到错误文件夹
                self._move_to_error_folder(file_path, error_folder)
                return False
            
        except Exception:
            # 将文件移动到错误文件夹
            self._move_to_error_folder(file_path, error_folder)
            return False
    
    def _find_academic_year_table(self, doc, year_suffix):
        """查找特定学年对应的学院意见表格"""
        # 先尝试查找包含"学年"的表格
        for i, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    try:
                        if year_suffix in cell.text and '学年' in cell.text:
                            # 找到学年标题后，检查当前表格和下一个表格
                            for next_idx in [i, i + 1]:
                                if next_idx < len(doc.tables):
                                    target_table = doc.tables[next_idx]
                                    if self._find_cell_by_text(target_table, "学院 意见")[0] is not None:
                                        return target_table
                    except Exception:
                        continue
        
        # 再尝试查找只包含年份的表格
        for i, table in enumerate(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    try:
                        cell_text = cell.text
                        if year_suffix in cell_text:
                            # 找到年份标题后，检查当前表格和下一个表格
                            for next_idx in [i, i + 1]:
                                if next_idx < len(doc.tables):
                                    target_table = doc.tables[next_idx]
                                    if self._find_cell_by_text(target_table, "学院 意见")[0] is not None:
                                        return target_table
                    except Exception:
                        continue
        
        return None
    
    def _find_cell_by_text(self, table, search_text):
        """在表格中查找包含指定文本的单元格"""
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                try:
                    cell_text = cell.text.replace('\n', ' ')
                    if search_text in cell_text:
                        return i, j, cell
                except Exception:
                    continue
        return None, None, None
    
    def _fill_evaluation_text(self, table, search_text, evaluation_text):
        """在表格中查找并填写评语"""
        try:
            row_idx, col_idx, cell = self._find_cell_by_text(table, search_text)
            if cell and row_idx is not None:
                # 填写评语到第二列（内容列）
                content_cell = table.rows[row_idx].cells[1] if len(table.rows[row_idx].cells) > 1 else cell
                content_cell.text = evaluation_text
                
                # 设置字体格式
                self._format_cell_text(content_cell)
                return True
        except Exception:
            pass
        return False
    
    def _fill_table_cell(self, table, row_idx, col_idx, text):
        """填写表格特定位置的单元格"""
        try:
            if len(table.rows) > row_idx and len(table.rows[row_idx].cells) > col_idx:
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = text
                self._format_cell_text(cell)
                return True
        except Exception:
            pass
        return False
    
    def _format_cell_text(self, cell):
        """设置单元格文本格式"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "宋体"
                run.font.size = Pt(10.5)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _move_to_error_folder(self, file_path, error_folder):
        """将文件移动到错误文件夹"""
        if not os.path.exists(error_folder):
            os.makedirs(error_folder)
        error_file_path = os.path.join(error_folder, os.path.basename(file_path))
        if os.path.exists(file_path):
            shutil.move(file_path, error_file_path)
    
    def process_class_files(self, class_dir: str) -> int:
        """
        处理某个班级文件夹中的所有文件
        
        Args:
            class_dir: 班级文件夹路径
            
        Returns:
            int: 成功处理的文件数量
        """
        if not os.path.exists(class_dir):
            print(f"❌ 班级文件夹不存在: {class_dir}")
            return 0
        
        # 获取所有docx文件
        docx_files = []
        for file in os.listdir(class_dir):
            if file.endswith('.docx') and not file.startswith('~'):
                docx_files.append(file)
        
        if not docx_files:
            print(f"⚠️  班级文件夹中未找到任何docx文件: {class_dir}")
            return 0
        print(f"📁 处理班级文件夹: {os.path.basename(class_dir)}")
        print(f"📄 找到 {len(docx_files)} 个文件需要处理")
        
        success_count = 0
        error_folder = os.path.join(class_dir, "处理失败的文件")
        
        for i, filename in enumerate(docx_files, 1):
            file_path = os.path.join(class_dir, filename)
            print(f"[{i}/{len(docx_files)}] 正在填写评语: {filename}")
            
            try:
                if self._process_single_file(file_path, error_folder):
                    success_count += 1
                    print(f"✓ 处理成功: {filename}")
                else:
                    print(f"✗ 处理失败: {filename}")
                    # _process_single_file 已经处理了错误文件的移动
                    
            except Exception as e:
                print(f"✗ 处理失败: {filename} - {str(e)}")
                self._move_to_error_folder(file_path, error_folder)
        
        print(f"📊 班级处理完成: 成功 {success_count}/{len(docx_files)} 个文件")
        
        # 如果错误文件夹为空，删除它
        if os.path.exists(error_folder) and not os.listdir(error_folder):
            os.rmdir(error_folder)
        
        return success_count
